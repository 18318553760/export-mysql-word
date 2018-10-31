import pymysql
from docx import Document
from docx.shared import Pt #字体设置
from docx.shared import Inches #字体设置
from docx.oxml.ns import qn #字体设置
connection=pymysql.connect(host='localhost',
                           user='root',
                           password='root',
                           db='information_schema',
                           port=3306,
                           charset='utf8')#information_schema 这个不能变，其他是你的数据库的连接信息
schema='artallex' #要导出的数据库的名字
cursor=connection.cursor()
sql="select table_name,table_comment from information_schema.tables where TABLE_SCHEMA = '"+str(schema)+"'"
cursor.execute(sql)
tableInfoList=cursor.fetchall()
doc=Document()
doc.styles['Normal'].font.name=u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
for tableInfo in tableInfoList:
    tableName=tableInfo[0]
    tableComment=tableInfo[1]
    table_explain = tableName+",注解："+tableComment+",对应数据库的表："
    tableInfoSql="SELECT C.COLUMN_NAME AS '字段名',C.COLUMN_TYPE AS '数据类型',C.IS_NULLABLE AS '允许为空',C.EXTRA AS 'PK',C.COLUMN_COMMENT AS '字段说明' FROM information_schema.COLUMNS C INNER JOIN TABLES T ON C.TABLE_SCHEMA = T.TABLE_SCHEMA AND C.TABLE_NAME = T.TABLE_NAME WHERE T.TABLE_SCHEMA = '"+str(schema)+"' and T.TABLE_NAME='"+str(tableName)+"'"
    cursor.execute(tableInfoSql)
    tableColumnInfoList = cursor.fetchall()
    p = doc.add_paragraph('')
    p.add_run(table_explain, style="Heading 1 Char")
    row=cursor.rowcount
    table=doc.add_table(rows=1,cols=5)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text='字段名'
    hdr_cells[1].text = '字段类型'
    hdr_cells[2].text = '允许为空'
    hdr_cells[3].text = 'PK'
    hdr_cells[4].text = '字段说明'
    for tableColumn in tableColumnInfoList:
        new_cells = table.add_row().cells
        new_cells[0].text=tableColumn[0]
        new_cells[1].text = tableColumn[1]
        new_cells[2].text = tableColumn[2]
        new_cells[3].text = tableColumn[3]
        new_cells[4].text = tableColumn[4]
    p = doc.add_paragraph('')
doc.save('I:/paoyou.docx')